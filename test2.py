from tkinter import *
import tkinter.ttk as ttk
from PIL import Image,ImageTk
import googletrans
from googletrans import Translator
from tkinter import filedialog
import os
import docx
import fitz
from tkinterdnd2 import *
import pyttsx3

#windowform
root=Tk()
root.title('PHIÊN DỊCH')
root.geometry('1000x600')
load=Image.open('D:\Python\\background.jpg')
render=ImageTk.PhotoImage(load)
img=Label(root,image=render)
img.place(x=0,y=0)

name=Label(root,text='PHIÊN DỊCH',fg='black',bd=0,bg='#FBFBFB')
name.config(font=('Helvetica',30))
name.pack(pady=50)

box1=Text(root,width=35,height=10,font=("ROBOTO",16))
box1.place(x=50,y=180)

box2=Text(root,width=35,height=10,font=("ROBOTO",16))
box2.place(x=520,y=180)
box2['state']='disabled'

language =[]

for x in googletrans.LANGUAGES.values():
    language.append(x.capitalize())


selected_language = StringVar()
combo = ttk.Combobox(root,values=language,font=('Helvetica',15),textvariable=selected_language)
#combo['values']=('Tiếng Việt','Tiếng Anh','Tiếng Nhật','Tiếng Hàn')
combo.place(x=30,y=140)
combo.current(0)

selected_language_2 = StringVar()
combo2 = ttk.Combobox(root,values=language,font=('Helvetica',15),textvariable=selected_language_2)
#combo['values']=('Tiếng Việt','Tiếng Anh','Tiếng Nhật','Tiếng Hàn')
combo2.place(x=500,y=140)
combo2.current(0)


Button_frame=Frame(root).pack(side=BOTTOM)
def clear():
    box2['state']='normal'
    box1.delete(1.0,END)
    box2.delete(1.0,END)
    box2['state']='disabled'


def translate():
    box2['state']='normal'
    box2.delete(1.0,END)
    INPUT = box1.get(1.0,END)
    lang_1=str(selected_language.get())
    lang_2=str(selected_language_2.get())
    t=googletrans.Translator()
    a= t.translate(INPUT, src=lang_1, dest=lang_2)
    b=a.text
    box2.insert(END,b)
    box2['state']='disabled'
    





# hàm xử lý file


def openFile():
    filepath = filedialog.askopenfilename(title="Chọn 1 file", filetypes=(
        ("text files", "*.txt"), ("word files", "*.docx"), ("pdf files", "*.pdf"), ("all files", "*.*")))
    basename = os.path.basename(filepath)
    name_tuple = os.path.splitext(basename)
    # xử lý file text
    if name_tuple[1] == '.txt':
        file = open(filepath, 'r')
        clear()
        box1.insert(END, file.read())
        file.close()
    # xử lý file word
    if name_tuple[1] == '.docx':
        doc = docx.Document(filepath)
        Text = []
        for para in doc.paragraphs:
            Text.append(para.text)
        clear()
        box1.insert(END, '\n'.join(Text))
    # xử lý file pdf
    if name_tuple[1] == '.pdf':
        doc = fitz.open(filepath)
        result = []
        for i in range(0, doc.page_count):
            page = doc.load_page(i)
            result = page.get_text()
        clear()
        box1.insert(END, ''.join(result))

# hàm kéo thả file


def DisplayText(event):
    filepath = event.data.strip('{}')
    basename = os.path.basename(filepath)
    name_tuple = os.path.splitext(basename)
    # xử lý file text
    if name_tuple[1] == '.txt':
        file = open(filepath, 'r')
        clear()
        box1.insert(END, file.read())
        file.close()
    # xử lý file word
    if name_tuple[1] == '.docx':
        doc = docx.Document(filepath)
        Text = []
        for para in doc.paragraphs:
            Text.append(para.text)
        clear()
        box1.insert(END, '\n'.join(Text))
    # xử lý file pdf
    if name_tuple[1] == '.pdf':
        doc = fitz.open(filepath)
        result = []
        for i in range(0, doc.page_count):
            page = doc.load_page(i)
            result = page.get_text()
        clear()
        box1.insert(END, ''.join(result))


# kéo thả file
box1.drop_target_register(DND_FILES)
box1.dnd_bind('<<Drop>>', DisplayText)

#Đọc file
Button_frame=Frame(root).pack(side=BOTTOM)
def clear():
    box1.delete(1.0,END)
    box2.delete(1.0,END)
def translate():
    INPUT = box1.get(1.0,END)
    print(INPUT)
    t = Translator()
    a = t.translate(INPUT, src='vi', dest='en')
    b=a.text 
    box2.insert(END,b)
def speach():
    SPEACH = box2.get(1.0,END)
    e = pyttsx3.init()
    e.say(SPEACH)
    e.runAndWait()


trans_button=Button(Button_frame,text='Dịch',height=2,width=10, font=(('Arial'),10,'bold'),bg='#303030',fg='#FFFFFF',command=translate)
trans_button.place(x=30,y=440) 
clear_button=Button(Button_frame,text='Xóa',height=2,width=10, font=(('Arial'),10,'bold'),bg='#303030',fg='#FFFFFF',command=clear)
clear_button.place(x=150,y=440)  
open_file_button = Button(Button_frame, text="Mở File", height=2, width=10, font=(('Arial'), 10, 'bold'), bg='#303030', fg='#FFFFFF', command=openFile)
open_file_button.place(x=270,y=440)
clear_button=Button(Button_frame,text='Đọc',height=2,width=10, font=(('Arial'),10,'bold'),bg='#303030',fg='#FFFFFF',command=speach)
clear_button.place(x=390,y=440)  

    
root.mainloop()
