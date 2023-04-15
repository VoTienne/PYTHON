import pyttsx3
from tkinter import *
from PIL import Image,ImageTk
from googletrans import Translator

# thư viện cần thêm vào
from pytesseract import pytesseract
import googletrans
import tkinter.ttk as ttk
from tkinter import filedialog
import os
import docx
import fitz
from tkinterdnd2 import *

#windowform
root=Tk()
root.title('PHIÊN DỊCH')
root.geometry('800x500')
load=Image.open('background.jpg')
render=ImageTk.PhotoImage(load)
img=Label(root,image=render)
img.place(x=0,y=0)

name=Label(root,text='PHIÊN DỊCH',fg='black',bd=0,bg='#FBFBFB')
name.config(font=('Helvetica',30))
name.pack(pady=10)

box1=Text(root,width=30,height=10,font=("ROBOTO",16))
box1.place(x=20,y=130)

box2=Text(root,width=30,height=10,font=("ROBOTO",16))
box2.place(x=420,y=130)
box2['state']='disabled'

#tạo list language và gán từng giá trị của googletrans vào list
language =[]
for x in googletrans.LANGUAGES.values():
    language.append(x.capitalize())

#combobox 1 chọn ngôn ngữ dịch
selected_language = StringVar()
combo = ttk.Combobox(root,values=language,state="readonly",font=('Helvetica',11),textvariable=selected_language)
combo.place(x=20,y=80)
combo.current(0)

#combobox 2 chọn ngôn ngữ cần dịch 
selected_language_2 = StringVar()
combo2 = ttk.Combobox(root,values=language,state="readonly",font=('Helvetica',11),textvariable=selected_language_2)
combo2.place(x=420,y=80)
combo2.current(0)

Button_frame=Frame(root).pack(side=BOTTOM)
def clear():
    box2['state']='normal'
    box1.delete(1.0,END)
    box2.delete(1.0,END)
    box2['state']='disabled'
def translate():
    box2['state']='normal'
    box2.delete(1.0,END)
    INPUT = box1.get(1.0,END)
    lang_1=str(selected_language.get())
    lang_2=str(selected_language_2.get())
    t=googletrans.Translator()
    a= t.translate(INPUT, src=lang_1, dest=lang_2)
    b=a.text
    box2.insert(END,b)
    box2['state']='disabled'
def speach():
    SPEACH = box2.get(1.0,END)
    e = pyttsx3.init()
    e.say(SPEACH)
    e.runAndWait()

# hàm xử lý file


def openFile():
    filepath = filedialog.askopenfilename(title="Chọn 1 file", filetypes=(
        ("text files", "*.txt"), ("word files", "*.docx"), ("pdf files", "*.pdf"), ("all files", "*.*")))
    basename = os.path.basename(filepath)
    name_tuple = os.path.splitext(basename)
    # xử lý file text
    if name_tuple[1] == '.txt':
        file = open(filepath, 'r')
        clear()
        box1.insert(END, file.read())
        file.close()
    # xử lý file word
    if name_tuple[1] == '.docx':
        doc = docx.Document(filepath)
        Text = []
        for para in doc.paragraphs:
            Text.append(para.text)
        clear()
        box1.insert(END, '\n'.join(Text))
    # xử lý file pdf
    if name_tuple[1] == '.pdf':
        doc = fitz.open(filepath)
        result = []
        for i in range(0, doc.page_count):
            page = doc.load_page(i)
            result = page.get_text()
        clear()
        box1.insert(END, ''.join(result))

# hàm kéo thả file


def DisplayText(event):
    filepath = event.data.strip('{}')
    basename = os.path.basename(filepath)
    name_tuple = os.path.splitext(basename)
    # xử lý file text
    if name_tuple[1] == '.txt':
        file = open(filepath, 'r')
        clear()
        box1.insert(END, file.read())
        file.close()
    # xử lý file word
    if name_tuple[1] == '.docx':
        doc = docx.Document(filepath)
        Text = []
        for para in doc.paragraphs:
            Text.append(para.text)
        clear()
        box1.insert(END, '\n'.join(Text))
    # xử lý file pdf
    if name_tuple[1] == '.pdf':
        doc = fitz.open(filepath)
        result = []
        for i in range(0, doc.page_count):
            page = doc.load_page(i)
            result = page.get_text()
        clear()
        box1.insert(END, ''.join(result))
    if (name_tuple[1] == '.jpg' or name_tuple[1] == '.jpeg' or name_tuple[1] == '.png'):
        path_to_tesseract = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        pytesseract.tesseract_cmd = path_to_tesseract
        img = Image.open(filepath)
        text = pytesseract.image_to_string(img,lang="eng")
        clear()
        box1.insert(END, text)
        print(text)


# kéo thả file
box1.drop_target_register(DND_FILES)
box1.dnd_bind('<<Drop>>', DisplayText)

trans_button=Button(Button_frame,text='Dịch',height=2,width=10, font=(('Arial'),10,'bold'),bg='#303030',fg='#FFFFFF',command=translate)
trans_button.place(x=20,y=400) 
clear_button=Button(Button_frame,text='Xóa',height=2,width=10, font=(('Arial'),10,'bold'),bg='#303030',fg='#FFFFFF',command=clear)
clear_button.place(x=140,y=400)
speach_button=Button(Button_frame,text='Đọc',height=2,width=10, font=(('Arial'),10,'bold'),bg='#303030',fg='#FFFFFF',command=speach)
speach_button.place(x=260,y=400)
open_file_button = Button(Button_frame, text="Mở File", height=2, width=10, font=(
('Arial'), 10, 'bold'), bg='#303030', fg='#FFFFFF', command=openFile)
open_file_button.place(x=380, y=400) 
    
root.mainloop()